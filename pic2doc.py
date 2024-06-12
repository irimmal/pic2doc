import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_COLOR_INDEX
import numpy as np
from  PIL import Image
import re
import colorsys
import argparse

def pic2doc(size:tuple, doc_path:str, pic_path:str, aspect_ratio=int(1), string='新岛夕天下第一', txt_method='random', font_type='宋体'):
    # 对于中文宋体，长宽比选择为1.45比较好
    # 建议使用等宽字体
    doc = docx.Document()
    if aspect_ratio != 1:
        size = (int(size[0]/aspect_ratio), size[1]) #size=纵向文字个数*横向文字个数。由于字体的长宽比不为1，因而进行缩放
    pic = pic2rgb(pic_path, size)
    txt = generate_txt(size, string, txt_method)
    para = doc.add_paragraph()
    para.width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER #设置段落居中对齐
    font_size = calc_font_size(size, doc)
    for i in range(size[0]*size[1]):
        # 逐个输入字符并设置颜色
        if i % size[1] == 0 and i != 0:
            para.add_run('\n')
        run = para.add_run(txt[i])
        rgb = pic[i]
        '''
        h,s,v = colorsys.rgb_to_hsv(*[x/255 for x in rgb])
        s *= 1.2
        s = min(s, 1)
        rgb = colorsys.hsv_to_rgb(h,s,v) *255
        '''
        run.font.name = font_type
        run.font.size = font_size
        run.font.bold = True
        run.font.color.rgb = docx.shared.RGBColor(int(rgb[0]),int(rgb[1]),int(rgb[2]))
        run.font.highlight_color = WD_COLOR_INDEX.AUTO
    run.add_break()
    doc.save(doc_path)


def pic2rgb(pic_path, size:tuple)->np.ndarray:
    # 将图片缩放后转化为RGB数组
    img = Image.open(pic_path)
    img = img.resize((size[1],size[0]))
    img.save(rewrite(pic_path, '_resized'))
    img = img.convert('RGB')
    img = np.reshape(img, (size[0]*size[1], 3))
    return img

def rewrite(filename, info)->str:
    # 使用正则表达式匹配最后一个点之前的部分
    match = re.match(r'^(.+)(\.[^.]+)$', filename)
    if match:
        name, extension = match.groups()
        # 在最后一个点之前插入信息
        new_filename = f"{name}{info}{extension}"
        return new_filename
    else:
        return None

def generate_txt(size:tuple, string='新岛夕天下第一', method='random')->list[str]:
    txt_len = size[0]*size[1]
    str_len = len(string)
    if method == 'repeat':
        txt = [string[i%str_len] for i in range(txt_len)]
    elif method == 'randint':
        txt = np.random.randint(0, 9, size)
    else:
        txt = [string[i%str_len] for i in np.random.randint(0,str_len,size=txt_len)]
    return txt

def calc_font_size(size:tuple, doc):
    # 计算字体大小。段落宽度/每行所需字符个数（横向像素个数）
    section = doc.sections[0]
    width = section.page_width - section.left_margin - section.right_margin
    return width/size[1]

def args_gene():
    parser = argparse.ArgumentParser()
    parser.add_argument('-s','--size', type=int, nargs=2, default=(100,100), help='size of the output document')
    parser.add_argument('-d','--doc_path', type=str, help='path of the output document')
    parser.add_argument('-p','--pic_path', type=str, help='path of the input picture')
    parser.add_argument('-ar','--aspect_ratio', type=float, default=1, help='aspect ratio of the font')
    parser.add_argument('-c','--string', type=str, default='新岛夕天下第一', help='characters to be used')
    parser.add_argument('-m','--txt_method', type=str, default='random', help='method to generate text')
    parser.add_argument('-f','--font_type', type=str, default='宋体', help='font type')
    args = parser.parse_args()
    return args

if __name__ == '__main__':
    args = args_gene()
    pic2doc(args.size, args.doc_path, args.pic_path, args.aspect_ratio, args.string, args.txt_method, args.font_type)